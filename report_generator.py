import os
import tempfile
import streamlit as st
import pdfkit
import openai
import json
from jinja2 import Template, Environment, BaseLoader
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from dotenv import load_dotenv
import shutil
from utils import add_hyperlink, process_links_for_template
from docx.shared import Pt

# Load environment variables from .env if present
load_dotenv()

# Constants
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/123.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Connection": "keep-alive"
}

def generate_seo_summary(data):
    """
    Generate an SEO analysis summary using OpenAI API.
    
    Args:
        data (dict): The SEO analysis data.
        
    Returns:
        str: The summary text or None if generation failed.
    """
    try:
        # Check if OpenAI API key is available in Streamlit secrets
        openai_api_key = st.secrets.get("OPENAI_API_KEY")
        if not openai_api_key:
            st.warning("OpenAI API key not found. Summary generation skipped.")
            return None
            
        # Create a client instance
        client = openai.OpenAI(api_key=openai_api_key)
        
        # Prepare the data to include in the prompt
        primary_keyword = data.get("keyword", "")
        serp_ranking = data.get("domain_organic_position", ">50")
        
        #Get Content Gap Analysis
        if data.get("content_gap_analysis") and data["content_gap_analysis"].get("results"):
            content_gap = data["content_gap_analysis"]["results"]
            
        # Get schema info
        schema_info = []
        if data.get("content_analysis", {}).get("schema_table"):
            for schema in data["content_analysis"]["schema_table"]:
                schema_info.append(f"{schema.get('schema', '')}: {schema.get('implemented', 'No')}")
                
        # Get image info
        images_info = []
        if data.get("content_analysis", {}).get("images"):
            for img in data["content_analysis"]["images"]:
                images_info.append(f"Source: {img.get('src', '')}, Alt: {img.get('alt', '')}")
        
        # Get YouTube info
        youtube_info = data.get("youtube_results", [])
        youtube_summary = []
        for yt in youtube_info[:3]:  # Limit to first 3 results
            youtube_summary.append(f"Title: {yt.get('title', '')}, Link: {yt.get('link', '')}")
            
        # Get 3rd party sites
        third_party_sites = []
        if data.get("popular_ai_overview_sites"):
            for site, urls in data["popular_ai_overview_sites"].items():
                for url in urls[:2]:  # Limit to first 2 URLs per site
                    third_party_sites.append(f"{site.capitalize()}: {url}")
        
        # Get social info
        social_info = []
        if data.get("social_channels"):
            for channel in data["social_channels"]:
                social_info.append(f"{channel.get('channel', '')}: {channel.get('suggestions', '')}")
        
        # Construct the prompt
        prompt = f"""Please summarize the SEO analysis for the given document, focusing on the following points:
1. **Optimize if SERP Ranking is [>20]**:
   * Highlight any keywords with SERP rankings above 20 and suggest optimization actions.
2. **Content-Specific Actions**:
   * Identify missing or underdeveloped headers.
   * Suggest content updates or additions to improve relevance and keyword alignment.
   * Recommend any new sections that competitors might have, which are missing from the current page.
3. **Schema to Implement**:
   * List schemas already implemented (if any) such as Breadcrumbs, FAQ, Article, Video, etc.
   * Identify any missing schemas that need to be added to improve content categorization, engagement, and SERP performance.
4. **YouTube Actions**:
   * Suggest relevant videos that should be embedded, and recommend optimizing video descriptions with timestamps and CTAs.
5. **Image Optimization**:
   * Indicate whether images are missing in key sections, especially between H1 and FAQ.
   * Suggest alt text optimization for any existing images.
6. **Opportunity on 3rd Party Sites**:
   * Provide a list of 3rd party sites (URLs) that may present backlink opportunities, and note if the client is mentioned on those sites.
7. **Posts on LinkedIn and Reddit**:
   * Suggest relevant LinkedIn posts or groups to join and engage with for improved visibility.
   * Recommend Reddit discussions or threads that could be leveraged to boost content engagement and authority.

Please base your answers solely on the content and data provided below, and present your suggestions as numbered bullet points for word document, do not give output in markdown format:

Primary Keyword: {primary_keyword}
SERP Ranking: {serp_ranking}

Content Gap Analysis: {content_gap}

Schema Information: {schema_info}

Images: {images_info}

YouTube: {youtube_summary}

Third-party Sites: {third_party_sites}

Social Channels: {social_info}

Content Analysis: {data.get('content_gap_analysis', {}).get('results', [])}

Format each point as a concise, actionable recommendation.
"""

        # Call the OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an SEO expert providing actionable recommendations based on data analysis."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        # Extract and return the summary text
        summary = response.choices[0].message.content.strip()
        return summary
        
    except Exception as e:
        st.error(f"Error generating SEO summary: {str(e)}")
        return None

def generate_docx_report(data, domain, output_file="aio_report.docx"):
    """
    Generate a Word document report based on the provided SEO analysis data.
    
    Args:
        data (dict): SEO analysis data.
        domain (str): Domain name.
        output_file (str, optional): Output file name. Defaults to "aio_report.docx".
        
    Returns:
        None
    """
    document = Document()
    document.add_heading("SEO Analysis Report", level=1)
    p = document.add_paragraph()
    p.add_run("Keyword: ").bold = True
    p.add_run(str(data.get("keyword", "")))
    p = document.add_paragraph()
    p.add_run("Target URL: ").bold = True
    p_link = document.add_paragraph()
    add_hyperlink(p_link, data.get("target_url", ""), data.get("target_url", ""))
    p = document.add_paragraph()
    p.add_run(domain) 
    p.add_run(" Found in AI Overview Sources: ").bold = True
    p.add_run(str(data.get("domain_found", "")))
    document.add_heading("Domain Ranking", level=2)
    # Create the table with header
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Keyword"
    hdr_cells[1].text = "Google Search"
    hdr_cells[2].text = "Google - AI Overview"

    # Add row for the primary keyword
    row_cells = table.add_row().cells
    row_cells[0].text = data.get("keyword", "")
    row_cells[1].text = str(data.get("domain_organic_position", "Not Ranking"))
    row_cells[2].text = str(data.get("domain_ai_position", "Not Ranking"))

    # Add rows for secondary keywords if they exist
    secondary_keywords = data.get("secondary_keywords", [])
    for kw in secondary_keywords:
        row_cells = table.add_row().cells
        row_cells[0].text = kw
        row_cells[1].text = str(data.get("domain_organic_position_secondary", {}).get(kw, "Not Ranking"))
        row_cells[2].text = str(data.get("domain_ai_position_secondary", {}).get(kw, "Not Ranking"))
    
    # Rest of the code for other sections
    document.add_heading("Other Pages from AI Overview Sources", level=3)

    if data.get("social_ai_overview_sites") and any(data["social_ai_overview_sites"].values()):
        document.add_heading("Social Sites:", level=4)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Site"
        hdr_cells[1].text = "URL"
        
        for site, urls in data.get("social_ai_overview_sites").items():
            for url in urls:
                row_cells = table.add_row().cells
                row_cells[0].text = site.capitalize()
                row_cells[1].text = f'\u2022 {url}'
    else:
        document.add_heading("No Social sites found on AI Overview", level=4)

    if data.get("popular_ai_overview_sites") and any(data["popular_ai_overview_sites"].values()):
        document.add_heading("3rd Party Popular Sites:", level=4)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Site"
        hdr_cells[1].text = "URL"
        
        for site, urls in data.get("popular_ai_overview_sites").items():
            for url in urls:
                row_cells = table.add_row().cells
                row_cells[0].text = site.capitalize()
                row_cells[1].text = f'\u2022 {url}'
    else:
        document.add_heading("No Popular sites found on AI Overview", level=4)

    if data.get("review_ai_overview_sites") and any(data["review_ai_overview_sites"].values()):
        document.add_heading("Review Sites:", level=4)
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Site"
        hdr_cells[1].text = "URL"
        
        for site, urls in data.get("review_ai_overview_sites").items():
            for url in urls:
                row_cells = table.add_row().cells
                row_cells[0].text = site.capitalize()
                row_cells[1].text = f'\u2022 {url}'
    else:
        document.add_heading("No Review sites found on AI Overview", level=4)

    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run("Number of AI Sources in Organic Search (first 20): ").bold = True
    p.add_run(str(data.get("ai_sources_in_organic_count", "")))

    # Add new Summary Section
    document.add_heading("SEO Analysis Summary: Actionable pain points", level=2)
    
    # Generate summary using OpenAI if API key is available
    summary = generate_seo_summary(data)
    
    if summary:
        # Split the summary into points and add them as numbered paragraphs
        points = summary.split("\n")
        for i, point in enumerate(points):
            if point.strip():  # Skip empty lines
                p = document.add_paragraph(style="List")
                p.add_run(point.strip())
    else:
        document.add_paragraph("Summary generation failed. Please check your OpenAI API key.")

    document.add_heading("Content Gap Analysis", level=2)
    
    # Check if we have gap analysis data
    if data.get("content_gap_analysis") and data["content_gap_analysis"].get("results"):
        # Create a table for content gap analysis
        gap_table = document.add_table(rows=1, cols=3)
        gap_table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = gap_table.rows[0].cells
        hdr_cells[0].text = "Category"
        hdr_cells[1].text = "Current Status"
        hdr_cells[2].text = "Suggestions"
        
        # Add data rows
        for result in data["content_gap_analysis"]["results"]:
            row_cells = gap_table.add_row().cells
            row_cells[0].text = result.get("category", "")
            row_cells[1].text = result.get("current_status", "")
            row_cells[2].text = result.get("suggestions", "")
    else:
        document.add_paragraph("No content gap analysis available.")
    
    document.add_heading("Images (After H1 and Before FAQ)", level=3)
    if data.get("content_analysis", {}).get("images"):
        tbl = document.add_table(rows=1, cols=3)  # Changed from 2 to 3 columns
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Image Source"
        hdr_cells[1].text = "Alt Text"
        hdr_cells[2].text = "Suggested Alt Text"  # New column
        
        for image in data["content_analysis"]["images"]:
            row_cells = tbl.add_row().cells
            row_cells[0].text = image.get("src", "")
            row_cells[1].text = image.get("alt", "")
            
            # Get suggested alt text if it's in the data
            if "suggested_alt" in image:
                row_cells[2].text = image.get("suggested_alt", "")
            else:
                row_cells[2].text = "No suggestion available"
    else:
        document.add_paragraph("No images found.")

    document.add_heading("Videos Embedded on Page", level = 3)
    if data.get("content_analysis", {}).get("videos"):
        tbl = document.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Video"
        hdr_cells[1].text = "Source"
        hdr_cells[2].text = "Remarks"
        for row in data["content_analysis"]["videos"]:
            row_cells = tbl.add_row().cells
            row_cells[0].text = row.get("tag", "")
            row_cells[1].text = row.get("src", "")
            row_cells[2].text = "It is good practice to add timestamps for key moments in the description."
    else:
        document.add_paragraph("No video found on page.")
    
    document.add_heading("Relevant Videos on Youtube Channel", level = 3)
    if data.get("relevant_video"):
        tbl = document.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Video"
        hdr_cells[1].text = "Suggestion"
        
        # Check if relevant_video is a list or a single item.
        relevant_video = data["relevant_video"]
        if isinstance(relevant_video, list):
            items = relevant_video
        else:
            # Wrap the non-list value in a list.
            items = [relevant_video]
        
        for item in items:
            row_cells = tbl.add_row().cells
            row_cells[0].text = str(item)
            row_cells[1].text = "Embed this most relevant video on page."
    else:
        document.add_paragraph("There is no relevant video on official channel.")
        document.add_paragraph("Create a video for this topic.")

    document.add_heading("Schema Markup", level=3)
    if data.get("content_analysis", {}).get("schema_table"):
        tbl = document.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Schema"
        hdr_cells[1].text = "Implemented"
        hdr_cells[2].text = "Remarks"
        for row in data["content_analysis"]["schema_table"]:
            row_cells = tbl.add_row().cells
            row_cells[0].text = row.get("schema", "")
            row_cells[1].text = row.get("implemented", "")
            row_cells[2].text = row.get("remarks", "")
    else:
        document.add_paragraph("No schema markup data found.")
    document.add_heading("Brand Mentions", level=2)
    document.add_heading("YouTube", level=3)
    if data.get("youtube_results"):
        tbl = document.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Title"
        hdr_cells[1].text = "Displayed Link"
        hdr_cells[2].text = "Source"
        hdr_cells[3].text = "Snippet"
        hdr_cells[4].text = "Key Moments"
        for yt in data["youtube_results"]:
            row_cells = tbl.add_row().cells
            p = row_cells[0].paragraphs[0]
            add_hyperlink(p, yt.get("link", ""), yt.get("title", ""))
            row_cells[1].text = yt.get("displayed_link", "")
            row_cells[2].text = yt.get("source", "")
            row_cells[3].text = yt.get("snippet", "")
            row_cells[4].text = yt.get("key_moments", "")
    else:
        document.add_paragraph("No YouTube results found.")
    document.add_heading("Suggestions:", level=4)
    for line in ["Upload videos frequently.", "Write keyword-rich descriptions with timestamps and CTAs."]:
         document.add_paragraph(line, style="List Bullet")
    
    document.add_heading("Social Channels", level=3)
    
    if data.get("social_channels"):
        tbl = document.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        
        # Table headers
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Social Channel"
        hdr_cells[1].text = "Relevant Articles / Questions"
        hdr_cells[2].text = "Suggestions"

        # Populate table rows
        for channel in data["social_channels"]:
            row_cells = tbl.add_row().cells
            row_cells[0].text = channel.get("channel", "")
            
            # Process multiple hyperlinks correctly
            relevant_text = channel.get("relevant", "")
            p = row_cells[1].paragraphs[0]
            p.style = "List Bullet"
            
            # Extracting links and titles properly
            if "<a href=" in relevant_text:
                import re
                links = re.findall(r"<a href='(.*?)' target='_blank'>(.*?)</a>", relevant_text)
                
                for idx, (url, title) in enumerate(links):
                    "<td>{{ item.keyword }}</td>"
                    add_hyperlink(p, url, title)
                    if idx < len(links) - 1:
                        p.add_run("\n\n")  # Add line break between links
            else:
                p.add_run(relevant_text)  # If no links, just add plain text
            
            row_cells[2].text = channel.get("suggestions", "")
    else:
        document.add_paragraph("No social channels data found.")

    # HORIZONTAL LINE ______________________________________________________________________________________________
    # Add a table to simulate a horizontal line
    table = document.add_table(rows=1, cols=1)

    # Access the cell and set its border properties
    cell = table.cell(0, 0)

    # Set the cell's border to create the appearance of a line
    cell.text = ""
    cell.paragraphs[0].runs[0].font.size = Pt(0)  # Set the font size to 0 to hide the text

    # You can use a cell border to simulate a horizontal line
    table.style = 'Table Grid'
    table.autofit = True    

    # HORIZONTAL LINE ______________________________________________________________________________________________

    document.add_heading("APPENDIX", level = 1).bold = True

    document.add_heading("All AI Overview Sources", level=2)
    
    # Get primary keyword AI Overview competitors
    ai_competitors = data.get("ai_overview_competitors", [])
    
    # Get AI Overview competitors from secondary keywords, if available
    secondary_ai_competitors = data.get("secondary_ai_overview_competitors", {})
    
    # Combine all AI Overview competitors
    all_ai_competitors = ai_competitors.copy()
    
    # Add sources from secondary keywords with keyword labeling
    for keyword, competitors in secondary_ai_competitors.items():
        for competitor in competitors:
            # Add keyword info to the competitor
            competitor_with_keyword = competitor.copy()
            competitor_with_keyword["keyword"] = keyword
            all_ai_competitors.append(competitor_with_keyword)
    
    # Remove duplicates based on URL
    unique_urls = set()
    unique_ai_competitors = []
    
    for competitor in all_ai_competitors:
        url = competitor.get("url")
        if url and url not in unique_urls:
            unique_urls.add(url)
            unique_ai_competitors.append(competitor)
    
    if unique_ai_competitors:
        # Create a table with 3 columns to include keyword information
        table = document.add_table(rows=1, cols=4)
        table.style = 'Light List'  # or use 'Table Grid' or any Word style you prefer

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'URL'
        hdr_cells[1].text = 'SERP Position'
        hdr_cells[2].text = 'AI Citation'
        hdr_cells[3].text = 'Source Keyword'

        for item in unique_ai_competitors:
            row_cells = table.add_row().cells
            url = item.get("url")
            # Use actual position if available, otherwise "> 50"
            position = item.get("position")
            position_text = str(position) if position is not None else "> 50"
            
            # Use actual citation if available, otherwise "Not Ranking"
            citation = item.get("citation")
            citation_text = str(citation) if citation is not None else "Not Ranking"
            
            keyword_source = item.get("keyword", data.get("keyword", "Primary"))  # Default to "Primary" for the main keyword

            # Add hyperlink to the first cell
            p = row_cells[0].paragraphs[0]
            add_hyperlink(p, url, url)

            # Add position to second cell
            row_cells[1].text = position_text
            
            # Add citation to third cell
            row_cells[2].text = citation_text
            # Add source keyword to fourth cell
            row_cells[3].text = keyword_source
    else:
        document.add_paragraph("No AI Overview Competitors found.")

    # Competitors Section
    document.add_heading("Competitors Listed", level=2)

    if "competitors" in data:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Competitor"
        hdr_cells[1].text = "AI Content"
        hdr_cells[2].text = "Source"

        for competitor in data["competitors"]:
            row_cells = table.add_row().cells
            row_cells[0].text = competitor.get("name", "")
            content = competitor.get("content", "")
            if isinstance(content, list):
                row_cells[1].text = "\n".join(f"• {item}" for item in content)  # Format as bullet list
            else:
                row_cells[1].text = content  # Keep as-is if not a list
            row_cells[2].text = competitor.get("source", "")

    document.add_heading("AI Overview Competitors Content Analysis", level=2)

    if data.get("aio_competitor_content"):
        for source, content in data["aio_competitor_content"].items():
            document.add_heading(source, level=3)

            # Images
            document.add_heading("Images", level=4)
            images = content.get("images", [])
            if images:
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Alt Text'
                hdr_cells[1].text = 'Image URL'

                for img in images:
                    row_cells = table.add_row().cells
                    row_cells[0].text = img.get("alt", "")
                    row_cells[1].text = img.get("src", "")
            else:
                document.add_paragraph("No images found.", style="BodyText")

            # Videos
            document.add_heading("Videos", level=4)
            videos = content.get("videos", [])
            if videos:
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Tag'
                hdr_cells[1].text = 'Video Source'

                for video in videos:
                    row_cells = table.add_row().cells
                    row_cells[0].text = video.get("tag", "").upper()
                    row_cells[1].text = video.get("src", "")
            else:
                document.add_paragraph("No videos found.", style="BodyText")

            # Schema Table
            document.add_heading("Schema Table", level=4)
            schema_table = content.get("schema_table", [])
            if schema_table:
                table = document.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Schema'
                hdr_cells[1].text = 'Implemented'

                for row in schema_table:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('schema', ''))
                    row_cells[1].text = str(row.get('implemented', ''))
            else:
                document.add_paragraph("No schema data found.", style="BodyText")
    else:
        document.add_paragraph("No citations data found.")

    document.add_heading(f'AI Overview Content for: {data.get("keyword")}', level=2)
    ai_content = data.get("ai_overview_content", "")
    for line in ai_content.split("\n"):
        document.add_paragraph(line)

    # Add secondary keyword AI overview content if available
    secondary_ai_content = data.get("secondary_ai_overview_content", {})
    
    # Get secondary keywords list safely
    secondary_keywords = data.get("secondary_keywords", [])
    
    # Only iterate if both secondary_keywords and secondary_ai_content are available
    if secondary_keywords and secondary_ai_content:
        for kw in secondary_keywords:
            if kw in secondary_ai_content:
                document.add_heading(f"\nAI Overview Content for: {kw}", level=2)
                ai_content_secondary = secondary_ai_content[kw]
                for line in ai_content_secondary.split("\n"):
                    document.add_paragraph(line)
            
    document.add_heading("Top SERP URLs", level=2)
    if data.get("competitor_urls"):
        for url in data["competitor_urls"]:
            p = document.add_paragraph(style="List Bullet")
            add_hyperlink(p, url, url)
    else:
        document.add_paragraph("No competitor URLs found.")

    document.add_heading("Headers", level=3)
    if data.get("content_analysis", {}).get("headers"):
        tbl = document.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        hdr_cells[0].text = "Header Tag"
        hdr_cells[1].text = "Header Text"
        for header in data["content_analysis"]["headers"]:
            row_cells = tbl.add_row().cells
            row_cells[0].text = header.get("tag", "")
            row_cells[1].text = header.get("text", "")
    else:
        document.add_paragraph("No headers found.")

    document.add_heading("Missing Headers (compared to AI Overview)", level=3)
    if data.get("content_analysis", {}).get("missing_headers"):
        document.add_paragraph(f"Missing Headers for keyword:: {data.get('keyword', 'Main Keyword')}:", style="List Bullet")
        for mh in data["content_analysis"]["missing_headers"]:
            document.add_paragraph(mh, style="List Bullet 2")
    else:
        document.add_paragraph(f"No missing headers compared to AI Overview for Target keyword: {data.get('keyword', 'Main Keyword')}.")
    
    # Secondary keywords missing headers - safely access data
    secondary_keywords = data.get("secondary_keywords", [])
    content_data_secondary = data.get("content_data_secondary", {})

    if secondary_keywords and content_data_secondary:
        for kw in secondary_keywords:
            if kw in content_data_secondary:
                missing = content_data_secondary[kw].get("missing_headers", [])
                if missing:
                    document.add_paragraph(f"Missing Headers for keyword:: {kw}:", style="List Bullet")
                    for mh in missing:
                        document.add_paragraph(mh, style="List Bullet 2")
                else:
                    document.add_paragraph(f"No missing headers compared to AI Overview for secondary keyword: {kw}.")

    document.save(output_file)
    st.success("DOCX report generated: " + output_file)